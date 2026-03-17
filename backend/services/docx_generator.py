import re
from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt


def _is_table_separator(line: str) -> bool:
    cleaned = line.strip()
    if "|" not in cleaned:
        return False
    return all(ch in "-:| " for ch in cleaned) and "-" in cleaned


def _is_table_row(line: str) -> bool:
    return "|" in line


def _parse_table_row(line: str) -> List[str]:
    parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [cell for cell in parts if cell != "" or len(parts) == 1]


def _add_code_block(doc: Document, lines: List[str]) -> None:
    for raw in lines:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(raw.rstrip("\n"))
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def generate_docx_from_markdown(markdown_text: str) -> bytes:
    doc = Document()
    lines = (markdown_text or "").splitlines()

    in_code = False
    code_lines: List[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if _is_table_row(line) and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _parse_table_row(line)
            rows: List[List[str]] = []
            j = i + 2
            while j < len(lines) and _is_table_row(lines[j]):
                rows.append(_parse_table_row(lines[j]))
                j += 1

            if header:
                col_count = len(header)
                table = doc.add_table(rows=1 + len(rows), cols=col_count)
                for col_idx, cell_text in enumerate(header):
                    cell = table.cell(0, col_idx)
                    cell.text = cell_text
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

                for row_idx, row in enumerate(rows, start=1):
                    padded = row + [""] * max(0, col_count - len(row))
                    for col_idx, cell_text in enumerate(padded[:col_count]):
                        table.cell(row_idx, col_idx).text = cell_text

            i = j
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            i += 1
            continue

        unordered_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if unordered_match:
            doc.add_paragraph(unordered_match.group(1).strip(), style="List Bullet")
            i += 1
            continue

        ordered_match = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if ordered_match:
            doc.add_paragraph(ordered_match.group(1).strip(), style="List Number")
            i += 1
            continue

        if line.strip() == "":
            doc.add_paragraph("")
            i += 1
            continue

        doc.add_paragraph(line.strip())
        i += 1

    if code_lines:
        _add_code_block(doc, code_lines)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
